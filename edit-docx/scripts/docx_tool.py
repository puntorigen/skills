#!/usr/bin/env python3
"""CLI tool for inspecting and editing .docx files while preserving formatting."""

from __future__ import annotations

import argparse
import copy
import re
import sys
from pathlib import Path
from typing import Optional

try:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph
except ImportError:
    print("ERROR: python-docx is not installed. Run: pip install python-docx", file=sys.stderr)
    sys.exit(1)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

BODY_TAG = qn("w:body")
PARA_TAG = qn("w:p")
TABLE_TAG = qn("w:tbl")
RUN_TAG = qn("w:r")
TEXT_TAG = qn("w:t")
RPR_TAG = qn("w:rPr")
DRAWING_TAG = qn("w:drawing")
PICT_TAG = qn("w:pict")
FIELD_CHAR_TAG = qn("w:fldChar")
FIELD_CODE_TAG = qn("w:instrText")


def _para_style_name(para: Paragraph) -> str:
    try:
        return para.style.name
    except Exception:
        return "Normal"


def _para_has_image(para: Paragraph) -> bool:
    for run in para.runs:
        r_elem = run._r
        if r_elem.findall(DRAWING_TAG) or r_elem.findall(PICT_TAG):
            return True
    return False


def _image_dims(para: Paragraph) -> Optional[str]:
    """Try to extract image dimensions from drawing extents (EMU -> inches)."""
    EMU_PER_INCH = 914400
    for run in para.runs:
        for drawing in run._r.findall(DRAWING_TAG):
            extents = drawing.findall(".//" + qn("wp:extent"))
            if extents:
                ext = extents[0]
                cx = int(ext.get("cx", 0))
                cy = int(ext.get("cy", 0))
                if cx and cy:
                    return f"{cx / EMU_PER_INCH:.1f}x{cy / EMU_PER_INCH:.1f}in"
    return None


def _body_elements(doc: Document):
    """Yield (tag, element) for each <w:p> and <w:tbl> in document order.

    Passes ``doc`` (not the raw body lxml element) as the parent so that
    style resolution via ``.part`` works correctly on the returned objects.
    """
    body = doc.element.body
    for child in body:
        if child.tag == PARA_TAG:
            yield ("para", Paragraph(child, doc))
        elif child.tag == TABLE_TAG:
            yield ("table", Table(child, doc))


def _cell_text(cell) -> str:
    return cell.text.strip().replace("\n", " ")


def _truncate(text: str, maxlen: int = 120) -> str:
    if len(text) <= maxlen:
        return text
    return text[: maxlen - 2] + ".."


# ---------------------------------------------------------------------------
# Run-aware text replacement engine
# ---------------------------------------------------------------------------

def _clone_rpr(run_element):
    """Clone the <w:rPr> from a run element, or return None."""
    rpr = run_element.find(RPR_TAG)
    if rpr is not None:
        return copy.deepcopy(rpr)
    return None


def _make_run_with_text(text: str, rpr_element=None):
    """Create a new <w:r> element with given text and optional formatting."""
    new_r = OxmlElement("w:r")
    if rpr_element is not None:
        new_r.append(copy.deepcopy(rpr_element))
    t = OxmlElement("w:t")
    t.text = text
    if text and (text[0] == " " or text[-1] == " "):
        t.set(qn("xml:space"), "preserve")
    new_r.append(t)
    return new_r


def _collect_run_map(paragraph: Paragraph):
    """Build a list of (run, run_element, start_offset, end_offset, text) for a paragraph."""
    mapping = []
    offset = 0
    for run in paragraph.runs:
        t = run.text or ""
        mapping.append((run, run._r, offset, offset + len(t), t))
        offset += len(t)
    return mapping


def _replace_in_paragraph(paragraph: Paragraph, find_text: str, replace_text: str, use_regex: bool = False) -> int:
    """Replace occurrences of find_text in a paragraph while preserving run formatting.

    Returns the number of replacements made.
    """
    full_text = paragraph.text
    if not full_text:
        return 0

    if use_regex:
        matches = list(re.finditer(find_text, full_text))
    else:
        matches = list(re.finditer(re.escape(find_text), full_text))

    if not matches:
        return 0

    run_map = _collect_run_map(paragraph)
    if not run_map:
        return 0

    for match in reversed(matches):
        m_start, m_end = match.start(), match.end()

        affected_runs = []
        for run, r_elem, r_start, r_end, r_text in run_map:
            if r_start < m_end and r_end > m_start:
                affected_runs.append((run, r_elem, r_start, r_end, r_text))

        if not affected_runs:
            continue

        first_run, first_r, first_start, first_end, first_text = affected_runs[0]
        rpr = _clone_rpr(first_r)

        prefix_in_first = first_text[: m_start - first_start]
        last_run, last_r, last_start, last_end, last_text = affected_runs[-1]
        suffix_in_last = last_text[m_end - last_start:]

        parent = first_r.getparent()

        replacement_r = _make_run_with_text(replace_text, rpr)
        first_r.addprevious(replacement_r)

        if prefix_in_first:
            prefix_r = _make_run_with_text(prefix_in_first, _clone_rpr(first_r))
            replacement_r.addprevious(prefix_r)

        if suffix_in_last:
            suffix_rpr = _clone_rpr(last_r)
            suffix_r = _make_run_with_text(suffix_in_last, suffix_rpr)
            replacement_r.addnext(suffix_r)

        for _, r_elem, _, _, _ in affected_runs:
            parent.remove(r_elem)

    paragraph._p  # force refresh
    final_runs = _collect_run_map(paragraph)
    return len(matches)


def _replace_in_paragraphs(paragraphs, find_text: str, replace_text: str, use_regex: bool = False) -> int:
    total = 0
    for para in paragraphs:
        total += _replace_in_paragraph(para, find_text, replace_text, use_regex)
    return total


# ---------------------------------------------------------------------------
# Commands
# ---------------------------------------------------------------------------

def cmd_inspect(args):
    doc = Document(args.file)

    sections = list(doc.sections)
    body_elems = list(_body_elements(doc))
    table_count = sum(1 for tag, _ in body_elems if tag == "table")
    elem_count = len(body_elems)

    print(f'=== DOCUMENT: {Path(args.file).name} ({len(sections)} section(s), {elem_count} element(s), {table_count} table(s)) ===')

    used_styles = set()
    for tag, elem in body_elems:
        if tag == "para":
            used_styles.add(_para_style_name(elem))
        elif tag == "table":
            for row in elem.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        used_styles.add(_para_style_name(p))
    if used_styles:
        print(f'Styles in use: {", ".join(sorted(used_styles))}')
    print()

    idx = 1
    table_num = 0

    print("--- BODY ---")
    for tag, elem in body_elems:
        if tag == "para":
            style = _para_style_name(elem)
            text = elem.text.strip()
            extras = ""
            if _para_has_image(elem):
                dims = _image_dims(elem)
                dim_str = f" {dims}" if dims else ""
                extras = f" [IMAGE{dim_str}]"
            print(f'[{idx}] {style}:{extras} "{_truncate(text)}"')
            idx += 1
        elif tag == "table":
            table_num += 1
            rows = elem.rows
            cols = len(rows[0].cells) if rows else 0
            print(f"[{idx}] TABLE #{table_num} ({len(rows)}x{cols}):")

            col_widths: list[int] = []
            cell_texts: list[list[str]] = []
            for row in rows:
                row_texts = [_truncate(_cell_text(c), 30) for c in row.cells]
                cell_texts.append(row_texts)
                for ci, ct in enumerate(row_texts):
                    while len(col_widths) <= ci:
                        col_widths.append(0)
                    col_widths[ci] = max(col_widths[ci], len(ct), 5)

            for ri, row_texts in enumerate(cell_texts):
                cells_str = " | ".join(ct.ljust(col_widths[ci]) for ci, ct in enumerate(row_texts))
                print(f"    Row {ri}: | {cells_str} |")

            idx += 1

    for si, section in enumerate(sections):
        header = section.header
        if header and header.paragraphs:
            has_content = any(p.text.strip() for p in header.paragraphs)
            if has_content:
                print(f"\n--- HEADER (Section {si + 1}) ---")
                for hi, p in enumerate(header.paragraphs):
                    text = p.text.strip()
                    if text:
                        print(f'[H{si + 1}.{hi + 1}] {_para_style_name(p)}: "{_truncate(text)}"')

        footer = section.footer
        if footer and footer.paragraphs:
            has_content = any(p.text.strip() for p in footer.paragraphs)
            if has_content:
                print(f"\n--- FOOTER (Section {si + 1}) ---")
                for fi, p in enumerate(footer.paragraphs):
                    text = p.text.strip()
                    if text:
                        print(f'[F{si + 1}.{fi + 1}] {_para_style_name(p)}: "{_truncate(text)}"')


def cmd_replace(args):
    doc = Document(args.file)
    total = 0

    if args.header or args.footer:
        for section in doc.sections:
            if args.header:
                total += _replace_in_paragraphs(
                    section.header.paragraphs, args.find, args.replace, args.regex
                )
            if args.footer:
                total += _replace_in_paragraphs(
                    section.footer.paragraphs, args.find, args.replace, args.regex
                )
    elif args.element:
        body_elems = list(_body_elements(doc))
        target_idx = args.element - 1
        if target_idx < 0 or target_idx >= len(body_elems):
            print(f"ERROR: Element index {args.element} out of range (1-{len(body_elems)})", file=sys.stderr)
            sys.exit(1)
        tag, elem = body_elems[target_idx]
        if tag == "para":
            total += _replace_in_paragraph(elem, args.find, args.replace, args.regex)
        elif tag == "table":
            for row in elem.rows:
                for cell in row.cells:
                    total += _replace_in_paragraphs(
                        cell.paragraphs, args.find, args.replace, args.regex
                    )
    else:
        body_elems = list(_body_elements(doc))
        for tag, elem in body_elems:
            if tag == "para":
                total += _replace_in_paragraph(elem, args.find, args.replace, args.regex)
            elif tag == "table":
                for row in elem.rows:
                    for cell in row.cells:
                        total += _replace_in_paragraphs(
                            cell.paragraphs, args.find, args.replace, args.regex
                        )

    out = args.output or args.file
    doc.save(out)
    print(f"Replaced {total} occurrence(s). Saved to {out}")


def cmd_insert(args):
    doc = Document(args.file)
    body_elems = list(_body_elements(doc))

    ref_idx = args.after - 1
    if ref_idx < 0 or ref_idx >= len(body_elems):
        print(f"ERROR: Element index {args.after} out of range (1-{len(body_elems)})", file=sys.stderr)
        sys.exit(1)

    _, ref_elem = body_elems[ref_idx]
    if hasattr(ref_elem, "_p"):
        ref_xml = ref_elem._p
    elif hasattr(ref_elem, "_tbl"):
        ref_xml = ref_elem._tbl
    else:
        print("ERROR: Cannot determine XML element for reference", file=sys.stderr)
        sys.exit(1)

    new_p = OxmlElement("w:p")
    ref_xml.addnext(new_p)
    new_para = Paragraph(new_p, doc)

    if args.text:
        new_para.add_run(args.text)
    if args.style:
        try:
            new_para.style = doc.styles[args.style]
        except KeyError:
            print(f"WARNING: Style '{args.style}' not found, using default", file=sys.stderr)

    out = args.output or args.file
    doc.save(out)
    print(f"Inserted paragraph after element [{args.after}]. Saved to {out}")


def cmd_edit_cell(args):
    doc = Document(args.file)
    body_elems = list(_body_elements(doc))

    tables = [(i + 1, elem) for i, (tag, elem) in enumerate(body_elems) if tag == "table"]
    if not tables:
        print("ERROR: No tables found in document", file=sys.stderr)
        sys.exit(1)

    target_table = None
    for tnum, (elem_idx, tbl) in enumerate(tables, 1):
        if tnum == args.table:
            target_table = tbl
            break

    if target_table is None:
        print(f"ERROR: Table #{args.table} not found (document has {len(tables)} table(s))", file=sys.stderr)
        sys.exit(1)

    rows = target_table.rows
    if args.row < 0 or args.row >= len(rows):
        print(f"ERROR: Row {args.row} out of range (0-{len(rows) - 1})", file=sys.stderr)
        sys.exit(1)

    cells = rows[args.row].cells
    if args.col < 0 or args.col >= len(cells):
        print(f"ERROR: Col {args.col} out of range (0-{len(cells) - 1})", file=sys.stderr)
        sys.exit(1)

    cell = cells[args.col]

    if args.append:
        if cell.paragraphs:
            cell.paragraphs[-1].add_run(args.text)
        else:
            cell.add_paragraph(args.text)
    else:
        if cell.paragraphs and cell.paragraphs[0].runs:
            first_run = cell.paragraphs[0].runs[0]
            rpr = _clone_rpr(first_run._r)
            for p in cell.paragraphs:
                for r in p.runs:
                    p._p.remove(r._r)
            new_r = _make_run_with_text(args.text, rpr)
            cell.paragraphs[0]._p.append(new_r)
        else:
            if cell.paragraphs:
                cell.paragraphs[0].text = args.text
            else:
                cell.add_paragraph(args.text)

    out = args.output or args.file
    doc.save(out)
    print(f"Edited table #{args.table} cell [{args.row},{args.col}]. Saved to {out}")


def cmd_add_row(args):
    doc = Document(args.file)
    body_elems = list(_body_elements(doc))

    tables = [(i + 1, elem) for i, (tag, elem) in enumerate(body_elems) if tag == "table"]
    if not tables:
        print("ERROR: No tables found in document", file=sys.stderr)
        sys.exit(1)

    target_table = None
    for tnum, (elem_idx, tbl) in enumerate(tables, 1):
        if tnum == args.table:
            target_table = tbl
            break

    if target_table is None:
        print(f"ERROR: Table #{args.table} not found (document has {len(tables)} table(s))", file=sys.stderr)
        sys.exit(1)

    values = [v.strip() for v in args.values.split(",")]
    expected_cols = len(target_table.rows[0].cells) if target_table.rows else 0

    new_row = target_table.add_row()

    if target_table.rows and len(target_table.rows) >= 2:
        template_row = target_table.rows[-2]
        for ci, cell in enumerate(new_row.cells):
            if ci < len(template_row.cells):
                tmpl_cell = template_row.cells[ci]
                if tmpl_cell.paragraphs and tmpl_cell.paragraphs[0].runs:
                    rpr = _clone_rpr(tmpl_cell.paragraphs[0].runs[0]._r)
                    text = values[ci] if ci < len(values) else ""
                    for p in cell.paragraphs:
                        for r in p.runs:
                            p._p.remove(r._r)
                    new_r = _make_run_with_text(text, rpr)
                    cell.paragraphs[0]._p.append(new_r)
                else:
                    cell.text = values[ci] if ci < len(values) else ""
            else:
                cell.text = values[ci] if ci < len(values) else ""
    else:
        for ci, cell in enumerate(new_row.cells):
            cell.text = values[ci] if ci < len(values) else ""

    out = args.output or args.file
    doc.save(out)
    print(f"Added row to table #{args.table} with {len(values)} value(s). Saved to {out}")


def cmd_delete(args):
    doc = Document(args.file)
    body_elems = list(_body_elements(doc))

    target_idx = args.element - 1
    if target_idx < 0 or target_idx >= len(body_elems):
        print(f"ERROR: Element index {args.element} out of range (1-{len(body_elems)})", file=sys.stderr)
        sys.exit(1)

    tag, elem = body_elems[target_idx]
    if tag == "para":
        elem._p.getparent().remove(elem._p)
    elif tag == "table":
        elem._tbl.getparent().remove(elem._tbl)

    out = args.output or args.file
    doc.save(out)
    print(f"Deleted element [{args.element}]. Saved to {out}")


def cmd_styles(args):
    doc = Document(args.file)

    para_styles = []
    char_styles = []
    table_styles = []

    for style in doc.styles:
        entry = f"  {style.name}"
        if style.type == WD_STYLE_TYPE.PARAGRAPH:
            para_styles.append(entry)
        elif style.type == WD_STYLE_TYPE.CHARACTER:
            char_styles.append(entry)
        elif style.type == WD_STYLE_TYPE.TABLE:
            table_styles.append(entry)

    print("=== PARAGRAPH STYLES ===")
    print("\n".join(sorted(para_styles)) if para_styles else "  (none)")
    print("\n=== CHARACTER STYLES ===")
    print("\n".join(sorted(char_styles)) if char_styles else "  (none)")
    print("\n=== TABLE STYLES ===")
    print("\n".join(sorted(table_styles)) if table_styles else "  (none)")


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(description="Inspect and edit .docx files preserving formatting.")
    sub = parser.add_subparsers(dest="command", required=True)

    # inspect
    p_inspect = sub.add_parser("inspect", help="Display document structure with element indices")
    p_inspect.add_argument("file", help="Path to .docx file")

    # replace
    p_replace = sub.add_parser("replace", help="Find/replace text preserving run formatting")
    p_replace.add_argument("file", help="Path to .docx file")
    p_replace.add_argument("--find", required=True, help="Text to find")
    p_replace.add_argument("--replace", required=True, help="Replacement text")
    p_replace.add_argument("--element", type=int, help="Only replace in element N (body index)")
    p_replace.add_argument("--header", action="store_true", help="Replace in headers")
    p_replace.add_argument("--footer", action="store_true", help="Replace in footers")
    p_replace.add_argument("--regex", action="store_true", help="Treat --find as a regex pattern")
    p_replace.add_argument("--output", "-o", help="Output file (default: overwrite input)")

    # insert
    p_insert = sub.add_parser("insert", help="Insert a paragraph after element N")
    p_insert.add_argument("file", help="Path to .docx file")
    p_insert.add_argument("--after", type=int, required=True, help="Element index to insert after")
    p_insert.add_argument("--text", default="", help="Paragraph text")
    p_insert.add_argument("--style", help="Paragraph style name (e.g. 'Heading 1', 'Normal')")
    p_insert.add_argument("--output", "-o", help="Output file (default: overwrite input)")

    # edit-cell
    p_ecell = sub.add_parser("edit-cell", help="Edit a specific table cell")
    p_ecell.add_argument("file", help="Path to .docx file")
    p_ecell.add_argument("--table", type=int, required=True, help="Table number (1-based)")
    p_ecell.add_argument("--row", type=int, required=True, help="Row index (0-based)")
    p_ecell.add_argument("--col", type=int, required=True, help="Column index (0-based)")
    p_ecell.add_argument("--text", required=True, help="New cell text")
    p_ecell.add_argument("--append", action="store_true", help="Append to cell instead of replacing")
    p_ecell.add_argument("--output", "-o", help="Output file (default: overwrite input)")

    # add-row
    p_arow = sub.add_parser("add-row", help="Add a row to a table")
    p_arow.add_argument("file", help="Path to .docx file")
    p_arow.add_argument("--table", type=int, required=True, help="Table number (1-based)")
    p_arow.add_argument("--values", required=True, help="Comma-separated cell values")
    p_arow.add_argument("--output", "-o", help="Output file (default: overwrite input)")

    # delete
    p_del = sub.add_parser("delete", help="Remove an element by body index")
    p_del.add_argument("file", help="Path to .docx file")
    p_del.add_argument("--element", type=int, required=True, help="Element index to delete (1-based)")
    p_del.add_argument("--output", "-o", help="Output file (default: overwrite input)")

    # styles
    p_styles = sub.add_parser("styles", help="List available document styles")
    p_styles.add_argument("file", help="Path to .docx file")

    args = parser.parse_args()

    dispatch = {
        "inspect": cmd_inspect,
        "replace": cmd_replace,
        "insert": cmd_insert,
        "edit-cell": cmd_edit_cell,
        "add-row": cmd_add_row,
        "delete": cmd_delete,
        "styles": cmd_styles,
    }
    dispatch[args.command](args)


if __name__ == "__main__":
    main()
